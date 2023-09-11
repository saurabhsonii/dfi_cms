from django.shortcuts import render, get_object_or_404
from django.contrib.auth import authenticate, login
from django.shortcuts import render, redirect
import datetime
from django.contrib import messages
from .forms import (LoginForm, ContactForm, AgentRegistrationForm,
                    AgentUpdateForm, LoanDetailsForm, PersonalDetailsForm, DocumentImagesForm,
                    VehicleDocumentsForm, DisbursementForm, OccupationDetailsForm, ChannalPatternForm, PatternSourcingForm)
from django.contrib.auth import logout
from .models import (Contact, CustomUser, OccupationDetails, DocumentImages,
                     LoanDetails, PersonalDetails, Disbursement, OccupationDetails, VehicleDocuments)
import os
from docx import Document
from django.http import HttpResponse, FileResponse
import html2text
from docx import Document

from django.contrib.auth.decorators import login_required

# Create your views here.


@login_required(login_url='login')
def index_view(request):
    contact_count = Contact.objects.all().count()
    agent_count = CustomUser.objects.filter(parent_id=request.user.id).count()
    vehicle_data = LoanDetails.objects.all().order_by("-created_at")

    context = {
        "contact_count": contact_count,
        "agent_count": agent_count,
        "vehicle_data": vehicle_data
    }
    return render(request, "dashboard/index.html", context)

# ---------------------------------contact section-------------------------------------------------------


# @login_required(login_url='login')
def contact_view(request):
    if request.method == 'POST':
        form = ContactForm(request.POST)
        if form.is_valid():
            form.save()
            messages.success(
                request, 'Your message has been sent successfully!')
            # Redirect to a success page or another view
            return redirect('contact')
        else:
            messages.error(
                request, 'There was an error in your submission. Please correct the errors below.')
    else:
        form = ContactForm()

    return render(request, 'dashboard/contact.html', {'form': form})


@login_required(login_url='login')
def contact_list(request):
    contacts = Contact.objects.all().order_by("-created_at")
    return render(request, 'dashboard/contact-list.html', {'contacts': contacts})


@login_required(login_url='login')
def contact_details(request, contact_id):
    contact = get_object_or_404(Contact, id=contact_id)
    contact.is_seen = True
    contact.save()
    return render(request, 'dashboard/contact-details.html', {'contact': contact})


@login_required(login_url='login')
def delete_contact(request, contact_id):
    contact = get_object_or_404(Contact, id=contact_id)
    contact.delete()
    return redirect('contact_list')
# -----------------------------------------------------------------------------------------------------------

# ----------------------------------------------------login section------------------------------------------


def login_view(request):
    if request.method == 'POST':
        form = LoginForm(request.POST)
        if form.is_valid():
            email = form.cleaned_data['email']
            password = form.cleaned_data['password']
            user = authenticate(request, username=email, password=password)
            if user is not None:
                login(request, user)
                # Change this to your desired URL after login
                return redirect('home')
            else:
                messages.error(
                    request, 'Invalid email or password. Please try again.')
    else:
        form = LoginForm()

    return render(request, 'dashboard/login.html', {'form': form})


@login_required(login_url='login')
def logout_view(request):
    logout(request)
    return redirect('home')
# -------------------------------------------------------------------------------------------------------------

# -----------------------------------------------------------------------------------


@login_required(login_url='login')
def register_agent(request):
    manager = request.user.id
    if request.method == 'POST':
        form = AgentRegistrationForm(request.POST)

        if form.is_valid():
            agent = form.save(commit=False)
            agent.parent_id = manager
            agent.user_role = 'agent'
            agent.is_agent = True
            agent.save()
            messages.success(request, 'Agent registered successfully.')
            # Redirect to the agent list page
            return redirect('agent_list')

    else:
        form = AgentRegistrationForm()

    return render(request, 'dashboard/register_agent.html', {'form': form})


@login_required(login_url='login')
def agent_list(request):
    manager = request.user.id
    agents = CustomUser.objects.filter(
        user_role='agent', parent_id=manager).order_by("-date_joined")
    return render(request, 'dashboard/agent-list.html', {'agents': agents})


@login_required(login_url='login')
def update_agent(request, agent_id):
    agent = get_object_or_404(CustomUser, id=agent_id, user_role='agent')

    if request.method == 'POST':
        form = AgentUpdateForm(request.POST, instance=agent)
        if form.is_valid():
            if (form.save()):
                messages.success(
                    request, 'Agent details updated successfully.')
            return redirect('agent_list')  # Redirect to the agent list page
    else:
        form = AgentUpdateForm(instance=agent)

    return render(request, 'dashboard/update-agent.html', {'form': form, 'agent': agent})


# -----------------------applicant ----------------------------------------------------------------------------
@login_required(login_url='login')
def applicantfrom(request):

    return render(request, "dashboard/applicant-form.html")

# /------------------------------------------vehicle--------------------------------


@login_required(login_url='login')
def vehicle_details(request):
    loan_type = request.GET.get('type')
    if request.method == 'POST':
        loan_type = request.GET.get('type')
        print(loan_type)
        form = LoanDetailsForm(request.POST)
        if form.is_valid():
            loan_data = form.cleaned_data
            loan = LoanDetails(**loan_data)
            loan.parent_id = request.user
            loan.loan_type = loan_type
            loan.save()
            loan_id = loan.id
            request.session['loan_data'] = loan_id
            messages.success(
                request, 'Your Vehicle details were saved successfully.')
            return redirect('personal_details')
    else:

        form = LoanDetailsForm()

    return render(request, 'dashboard/applicant-form.html', {'form': form, "loan_type": loan_type})


@login_required(login_url='login')
def personal_details(request):
    if request.method == 'POST':
        form = PersonalDetailsForm(request.POST)
        loan_data = request.session.get('loan_data', {})
        loan_data = LoanDetails.objects.get(id=loan_data)

        if form.is_valid():
            personal_data = form.cleaned_data
            personal = PersonalDetails(**personal_data)
            personal.vehicle_id = loan_data
            personal.save()
            personal = personal.id

            # Serialize the 'state' field to a dictionary
            # state_instance = personal_data['state']
            # state_dict = {
            #     'id': state_instance.id,
            #     'state_name': state_instance.state_name,
            #     'state_code': state_instance.state_code,
            # }
            # personal_data['state'] = state_dict
            request.session['personal_data'] = personal
            messages.success(
                request, 'Your Personal details were saved successfully.')
            return redirect('occupation_details')
    else:
        form = PersonalDetailsForm()

    return render(request, 'dashboard/personal-detail-step2.html', {'form': form, 'step': 2})


@login_required(login_url='login')
def occupation_details(request):
    if request.method == 'POST':
        form = DocumentImagesForm(request.POST, request.FILES)
        form2 = OccupationDetailsForm(request.POST, request.FILES)

        if form.is_valid() and form2.is_valid():

            # Save the uploaded images to your media directory
            uploaded_images = []
            for uploaded_file in request.FILES.getlist('image'):
                document = DocumentImages(
                    name=uploaded_file.name, image=uploaded_file)
                document.save()
                uploaded_images.append(document.image.path)

            # Store the file paths in the session
            # request.session['uploaded_image_paths'] = uploaded_images

            # uploaded_image_paths = request.session.get(
            #     'uploaded_image_paths', [])
            personal_data = request.session.get('personal_data')
            personal_data = PersonalDetails.objects.get(id=personal_data)
            applicant_id = personal_data.id

            occupation_data = form2.cleaned_data
            occupation = OccupationDetails(**occupation_data)
            occupation.applicant = personal_data
            occupation.save()

            # occupation_data = OccupationDetails.objects.create(
            #     applicant=applicant_id,
            #     role=form2.cleaned_data['role'],
            #     bank_statement=form2.cleaned_data['bank_statement'],
            #     salary_slip=form2.cleaned_data['salary_slip'],
            #     shop_act=form2.cleaned_data['shop_act'],
            #     itr=form2.cleaned_data['itr'],
            #     id_card=form2.cleaned_data['id_card'],
            #     offer_later=form2.cleaned_data['offer_later']
            # )

            # Associate uploaded images with the OccupationDetails instance
            for image_path in uploaded_images:
                document = DocumentImages(
                    name=image_path.split('/')[-1], image=image_path)
                document.save()
                occupation.document_image.add(document)
            messages.success(
                request, 'Your Occupation details were saved successfully.')
            return redirect('vehicle_documents')
    else:
        form = DocumentImagesForm()
        form2 = OccupationDetailsForm()

    return render(request, 'dashboard/occupentional-detail-step.html', {'form': form, 'form2': form2})


@login_required(login_url='login')
def vehicle_documents(request):
    if request.method == 'POST':
        vehicle_documents_form = VehicleDocumentsForm(
            request.POST, request.FILES)
        vehicle_data = request.session.get('loan_data', {})
        print(vehicle_data)
        vehicle_data = LoanDetails.objects.get(id=vehicle_data)

        if vehicle_documents_form.is_valid():
            # Handle the VehicleDocuments uploads
            vehicle_documents = vehicle_documents_form.save(commit=False)
            vehicle_documents.applicant = vehicle_data
            messages.success(
                request, 'Your Vehicle Document Uploaded saved successfully.')
            vehicle_documents.save()

            # Redirect to confirmation page or another appropriate page
            return redirect('ChannalPattern')

    else:
        form = VehicleDocumentsForm()

    return render(request, 'dashboard/vehicle-documets-step4.html', {'form': form})


def ChannalPattern(request):
    if request.method == 'POST':
        ChannalPattern_form = ChannalPatternForm(request.POST)
        vehicle_data = request.session.get('loan_data', {})
        vehicle_data = LoanDetails.objects.get(id=vehicle_data)

        if ChannalPattern_form.is_valid():
            ChannalPattern_data = ChannalPattern_form.save(commit=False)
            ChannalPattern_data.vehicle_id = vehicle_data

            messages.success(
                request, 'Your ChannalPattern details were saved successfully.')
            ChannalPattern_data.save()
            print(ChannalPattern_data.sourcing)
            if ChannalPattern_data.sourcing == "direct":
                return redirect('disbursement')
            else:
                return redirect('PatternSourcing')

    else:
        form = ChannalPatternForm()

    return render(request, 'dashboard/channalpattern-form-step5.html', {'form': form})


def PatternSourcing(request):
    if request.method == 'POST':
        PatternSourcing_Form = PatternSourcingForm(request.POST)
        vehicle_data = request.session.get('loan_data', {})
        vehicle_data = LoanDetails.objects.get(id=vehicle_data)

        if PatternSourcing_Form.is_valid():
            PatternSourcing_data = PatternSourcing_Form.save(commit=False)
            PatternSourcing_data.vehicle_id = vehicle_data

            messages.success(
                request, 'Your PatternSourcing details were saved successfully.')
            PatternSourcing_data.save()
            return redirect('disbursement')

    else:
        form = PatternSourcingForm()

    return render(request, 'dashboard/PatternSourcing-form.html', {'form': form})


@login_required(login_url='login')
def disbursement(request):
    if request.method == 'POST':
        disbursement_form = DisbursementForm(request.POST, request.FILES)
        vehicle_data = request.session.get('loan_data', {})
        vehicle_data = LoanDetails.objects.get(id=vehicle_data)

        if disbursement_form.is_valid():
            # Process and save the Disbursement data
            disbursement_data = disbursement_form.save(commit=False)
            disbursement_data.vehicle_id = vehicle_data
            # Associate the Disbursement data with the relevant VehicleDetails or OccupationDetails, if needed
            messages.success(
                request, 'Your Disbursment details were saved successfully.')
            disbursement_data.save()
            return redirect('applicants')

    else:
        form = DisbursementForm()

    return render(request, 'dashboard/disbursement-form-step6.html', {'form': form})


@login_required(login_url='login')
def confirmation(request):
    vehicle_data = request.session.get('vehicle_data', {})
    personal_data = request.session.get('personal_data', {})

    # You can save the data to the database here if needed
    vehicle = LoanDetails(**vehicle_data)
    vehicle.parent_id = request.user
    vehicle.save()

    # Deserialize the 'state' field from the dictionary
    # state_data = personal_data.get('state', {})

    # # Retrieve the 'State' instance based on the data in the dictionary
    # state_instance = State.objects.get(id=state_data['id'])

    # Create a new 'PersonalDetails' instance with the retrieved 'State' instance
    personal_details = PersonalDetails(
        title=personal_data['title'],
        applicant_name=personal_data['applicant_name'],
        father_name=personal_data['father_name'],
        mother_name=personal_data['mother_name'],
        applicant_email=personal_data['applicant_email'],
        contact=personal_data['contact'],
        address=personal_data['address'],
        city=personal_data['city'],
        state=personal_data['state'],  # Assign the 'State' instance
        county=personal_data['county'],
        pincode=personal_data['pincode'],
        vehicle_id=vehicle  # Assign the 'VehicleDetails' instance
    )

    # Save the 'PersonalDetails' instance to the database
    personal_details.save()

    uploaded_image_paths = request.session.get('uploaded_image_paths', [])
    applicant_id = personal_details.id

    occupation_data = OccupationDetails.objects.create(
        applicant_id=applicant_id)

    # Associate uploaded images with the OccupationDetails instance
    for image_path in uploaded_image_paths:
        document = DocumentImages(
            name=image_path.split('/')[-1], image=image_path)
        document.save()
        occupation_data.document_image.add(document)

    # Clear the session variable after processing the data
    request.session['uploaded_image_paths'] = []

    return render(request, 'dashboard/index.html', {'vehicle_data': vehicle_data, 'personal_data': personal_data})


@login_required(login_url='login')
def ApplicantView(request):
    vehicle_data = LoanDetails.objects.all().order_by("-created_at")

    return render(request, "dashboard/applicant-list.html", {"vehicle_data": vehicle_data})


@login_required(login_url='login')
def edit_vehicle(request, vehicle_id):
    vehicle = get_object_or_404(LoanDetails, id=vehicle_id)
    personal_data = get_object_or_404(PersonalDetails, vehicle_id=vehicle_id)
    occupation_data = get_object_or_404(
        OccupationDetails, applicant=personal_data)
    disbursement_data = get_object_or_404(Disbursement, vehicle_id=vehicle_id)
    VehicleDocuments_data = get_object_or_404(
        VehicleDocuments, applicant=vehicle_id)

    # Retrieve the associated DocumentImages for the OccupationDetails instance
    document_images = occupation_data.document_image.all()
    # get creater name who filled the form
    creater = get_object_or_404(CustomUser, id=vehicle.parent_id.id)

    if request.method == 'POST':
        vehicle_form = LoanDetailsForm(request.POST, instance=vehicle)
        personal_form = PersonalDetailsForm(
            request.POST, instance=personal_data)
        occupation_form = OccupationDetailsForm(
            request.POST, instance=occupation_data)
        disbursement_form = DisbursementForm(
            request.POST, instance=disbursement_data)
        VehicleDocuments_form = VehicleDocumentsForm(
            request.POST, request.FILES, instance=VehicleDocuments_data)

        document_form = DocumentImagesForm(request.POST, request.FILES)

        if (
            vehicle_form.is_valid() and
            personal_form.is_valid() and
            occupation_form.is_valid() and
            disbursement_form.is_valid() and
            VehicleDocuments_form.is_valid()
        ):
            vehicle_form.save()
            personal_form.save()
            occupation_form.save()
            disbursement_form.save()
            VehicleDocuments_form.save()

            # Handle the document uploads
            for uploaded_file in request.FILES.getlist('image'):
                document = DocumentImages(
                    name=uploaded_file.name, image=uploaded_file)
                document.save()
                occupation_data.document_image.add(document)

            # Redirect to the confirmation page after updating the data
            return redirect('applicants')

    else:
        vehicle_form = LoanDetailsForm(instance=vehicle)
        personal_form = PersonalDetailsForm(instance=personal_data)
        occupation_form = OccupationDetailsForm(instance=occupation_data)
        disbursement_form = DisbursementForm(instance=disbursement_data)
        VehicleDocuments_form = VehicleDocumentsForm(
            instance=VehicleDocuments_data)
        document_form = DocumentImagesForm()

    return render(request, 'dashboard/applicant-details.html', {
        'vehicle_form': vehicle_form,
        'personal_form': personal_form,
        'occupation_form': occupation_form,
        'disbursement_form': disbursement_form,
        'vehicleDocuments_form': VehicleDocuments_form,
        'document_form': document_form,
        'document_images': document_images,
        'creater': creater
    })


# -----------------------Home Loan-----------
@login_required(login_url='login')
def home_details(request):
    loan_type = request.GET.get('type')
    if request.method == 'POST':
        loan_type = request.GET.get('type')
        print(loan_type)
        form = LoanDetailsForm(request.POST)
        if form.is_valid():
            loan_data = form.cleaned_data
            loan = LoanDetails(**loan_data)
            loan.parent_id = request.user
            loan.loan_type = loan_type
            messages.success(request, 'Your details were saved successfully.')
            loan.save()
            loan_id = loan.id
            request.session['loan_data'] = loan_id

            return redirect('personal_details')
    else:

        form = LoanDetailsForm()

    return render(request, 'dashboard/home-form.html', {'form': form, "loan_type": loan_type})

# -----------------------Business Loan-----------


@login_required(login_url='login')
def business_details(request):
    loan_type = request.GET.get('type')
    if request.method == 'POST':
        loan_type = request.GET.get('type')
        print(loan_type)
        form = LoanDetailsForm(request.POST)
        if form.is_valid():
            loan_data = form.cleaned_data
            loan = LoanDetails(**loan_data)
            loan.parent_id = request.user
            loan.loan_type = loan_type
            messages.success(request, 'Your details were saved successfully.')
            loan.save()
            loan_id = loan.id
            request.session['loan_data'] = loan_id

            return redirect('personal_details')
    else:

        form = LoanDetailsForm()

    return render(request, 'dashboard/business-form.html', {'form': form, "loan_type": loan_type})


# -----------------------Micro Loan-----------
@login_required(login_url='login')
def micro_details(request):
    loan_type = request.GET.get('type')
    if request.method == 'POST':
        loan_type = request.GET.get('type')
        form = LoanDetailsForm(request.POST)
        if form.is_valid():
            loan_data = form.cleaned_data
            loan = LoanDetails(**loan_data)
            loan.parent_id = request.user
            loan.loan_type = loan_type
            messages.success(request, 'Your details were saved successfully.')
            loan.save()
            loan_id = loan.id
            request.session['loan_data'] = loan_id

            return redirect('personal_details')
    else:

        form = LoanDetailsForm()

    return render(request, 'dashboard/micro-form.html', {'form': form, "loan_type": loan_type})

# -----------------------Gold Loan-----------


@login_required(login_url='login')
def gold_details(request):
    loan_type = request.GET.get('type')
    if request.method == 'POST':
        loan_type = request.GET.get('type')
        print(loan_type)
        form = LoanDetailsForm(request.POST)
        if form.is_valid():
            loan_data = form.cleaned_data
            loan = LoanDetails(**loan_data)
            loan.parent_id = request.user
            loan.loan_type = loan_type
            messages.success(request, 'Your details were saved successfully.')
            loan.save()
            loan_id = loan.id
            request.session['loan_data'] = loan_id

            return redirect('personal_details')
    else:

        form = LoanDetailsForm()

    return render(request, 'dashboard/gold-loan-form.html', {'form': form, "loan_type": loan_type})


def approve(request, disburs_id):
    personal_form = Disbursement.objects.get(id=disburs_id)
    personal_form.status = False
    if (personal_form.save()):
        messages.success(request, 'Your Loan Application Update successfully.')
    return redirect('applicants')


def generate_loan_details_docx(request, vehicle_id):
    output_path = 'loan_details.docx'
    # Retrieve LoanDetails for the specified user
    vehicle_data = LoanDetails.objects.all().order_by("-created_at")
    vehicle = get_object_or_404(LoanDetails, id=vehicle_id)
    personal_data = get_object_or_404(PersonalDetails, vehicle_id=vehicle_id)
    occupation_data = get_object_or_404(
        OccupationDetails, applicant=personal_data)
    disbursement_data = get_object_or_404(Disbursement, vehicle_id=vehicle_id)
    VehicleDocuments_data = get_object_or_404(
        VehicleDocuments, applicant=vehicle_id)
    vehicle_form = LoanDetailsForm(instance=vehicle)
    personal_form = PersonalDetailsForm(instance=personal_data)
    occupation_form = OccupationDetailsForm(instance=occupation_data)
    disbursement_form = DisbursementForm(instance=disbursement_data)
    VehicleDocuments_form = VehicleDocumentsForm(
        instance=VehicleDocuments_data)
    document_form = DocumentImagesForm()
    # Retrieve the associated DocumentImages for the OccupationDetails instance
    document_images = occupation_data.document_image.all()
    # get creater name who filled the form
    creater = get_object_or_404(CustomUser, id=vehicle.parent_id.id)

    # Create a new Word document
    doc = Document()

    # Add a title to the document
    doc.add_heading('Loan Details Report', 0)

    # Loop through loan details and add them to the document
    for vehicle in vehicle_data:
        doc.add_heading(f'Loan Type: {vehicle.loan_type}', level=1)
        doc.add_paragraph(f'Vehicle Name: {vehicle.vehicle_name}')
        doc.add_paragraph(f'Vehicle Model: {vehicle.vehicle_model}')
        doc.add_paragraph(f'Vehicle Number: {vehicle.vehicle_number}')
        # Add more loan detail fields as needed
    docx_content = """
<html>
<head>
    <title>Loan Details</title>
</head>
<body>
    <h1>Loan Type: Car Loan</h1>
    <p>Vehicle Name: My Car</p>
    <p>Vehicle Model: XYZ123</p>
    <p>Vehicle Number: ABC123</p>
    <!-- Add more HTML content here -->
</body>
</html>
"""
    # Save the document to the specified output path
    doc.save(output_path)
    response = HttpResponse(
        docx_content, content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = f'attachment; filename="loan_details.docx"'

    return response
# Example usage:
# Replace 'user_id' with the actual user ID and 'output_path' with the desired file path

# @login_required(login_url='login')
# def approved(request, loan_id):
#     disbursement = get_object_or_404(Disbursement, id=loan_id)


#     return render(request, 'dashboard/update-agent.html', {'form': form, 'agent': agent})
